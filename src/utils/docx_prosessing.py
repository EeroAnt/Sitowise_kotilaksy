import io
import re
import streamlit as st
from docx import Document
from src.utils.field_mappings import get_unique_keys
from streamlit.runtime.uploaded_file_manager import UploadedFile

def process_doc(uploaded_file: UploadedFile) -> None:
    doc = Document(uploaded_file)
    fields_to_fill = []

    # We're assuming that the placeholders are nice enough
    # so this is enough parsing
    
    pattern = r"\{\{(\w+)\}\}"

    # Tables

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(pattern, para.text)
                    fields_to_fill.extend(matches)

    # Paragraphs
    
    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        fields_to_fill.extend(matches)
    
    st.session_state["field_mappings"] = get_unique_keys(fields_to_fill)

    return

def create_filled_doc(uploaded_file: UploadedFile, field_mappings: list[tuple[str, str]]) -> io.BytesIO:
    doc = Document(uploaded_file)
    
    mapping_iter = iter(field_mappings)
    pattern = r"\{\{\w+\}\}"
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    while re.search(pattern, para.text):
                        original, unique_key = next(mapping_iter)
                        placeholder = f"{{{{{original}}}}}"
                        value = st.session_state.get(unique_key, "")
                        para.text = para.text.replace(placeholder, value, 1)

    # Paragraphs

    for para in doc.paragraphs:
        while re.search(pattern, para.text):
            original, unique_key = next(mapping_iter)
            placeholder = f"{{{{{original}}}}}"
            value = st.session_state.get(unique_key, "")
            para.text = para.text.replace(placeholder, value, 1)
    

    # Save to BytesIO for download
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer